from datetime import datetime
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE

OUT = "Delta-NIDS-Codebase-Analysis-Report.docx"

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)

def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(str(text))
    r.bold = bold
    if color:
        r.font.color.rgb = RGBColor(*color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        set_cell_text(t.rows[0].cells[i], h, True, (255,255,255))
        shade(t.rows[0].cells[i], '17365D')
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            if len(t.rows) % 2 == 0:
                shade(cells[i], 'EAF2F8')
    if widths:
        for row in t.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Inches(width)
    doc.add_paragraph()
    return t

def bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet' if level == 0 else 'List Bullet 2')
    p.add_run(text)
    return p

def numbered(doc, text):
    p = doc.add_paragraph(style='List Number')
    p.add_run(text)
    return p

def heading(doc, text, level=1):
    doc.add_heading(text, level=level)

def paragraph(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        p.add_run(bold_prefix).bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p

doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(0.7)
sec.bottom_margin = Inches(0.7)
sec.left_margin = Inches(0.8)
sec.right_margin = Inches(0.8)

styles = doc.styles
styles['Normal'].font.name = 'Aptos'
styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Aptos')
styles['Normal'].font.size = Pt(10)
for name, size, color in [('Title', 30, '17365D'), ('Heading 1', 18, '17365D'), ('Heading 2', 13, '2F75B5'), ('Heading 3', 11, '404040')]:
    styles[name].font.name = 'Aptos Display'
    styles[name].font.size = Pt(size)
    styles[name].font.color.rgb = RGBColor.from_string(color)

# Header/footer
header = sec.header.paragraphs[0]
header.text = 'Delta-NIDS  |  Engineering Codebase Analysis'
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
header.runs[0].font.size = Pt(8)
header.runs[0].font.color.rgb = RGBColor(100,100,100)
footer = sec.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.add_run('Confidential engineering review  •  Generated 30 August 2026').font.size = Pt(8)

# Cover
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('\n\n')
r = p.add_run('DELTA-NIDS')
r.bold = True; r.font.size = Pt(34); r.font.color.rgb = RGBColor(23,54,93)
p.add_run('\n')
r = p.add_run('Full Codebase Analysis Report')
r.bold = True; r.font.size = Pt(22); r.font.color.rgb = RGBColor(47,117,181)
p.add_run('\n\n')
r = p.add_run('Architecture • Security • Quality • Testing • Recommendations')
r.italic = True; r.font.size = Pt(12)
p.add_run('\n\n\n')
r = p.add_run('Repository: DavidKuriyan/delta-ids\nBranch reviewed: main\nReview date: 30 August 2026\nHEAD: 0b9b31e')
r.font.size = Pt(11)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

heading(doc, 'Executive Summary', 1)
paragraph(doc, 'Delta-NIDS is a passive, cross-platform Network Intrusion Detection System combining a C++17 native engine with a Python/Scapy capture and persistence path, a native HTTP API, SQLite storage, and a Flask dashboard. The implementation is structured around a clear evidence flow: capture and normalization → packet decoding → flow tracking and protocol identification → signature and behavioral detection → alert/incident persistence → API and dashboard presentation.')
paragraph(doc, 'Overall assessment: the codebase is in a strong prototype-to-early-production foundation state. The repository demonstrates thoughtful defensive behavior, meaningful separation of concerns, and unusually broad deterministic tests for a small IDS. The principal release risks are operational rather than fundamental: CI is not represented in the repository, fuzzing/sanitizer coverage is absent, the Python and C++ paths duplicate some responsibilities, the dashboard is explicitly development-oriented, and no license is declared.')

table(doc, ['Area', 'Assessment', 'Rationale'], [
    ('Architecture', 'Strong', 'Layered C++ core with platform boundary and shared SQLite/API contract.'),
    ('Security posture', 'Good foundation', 'Bounds checks, passive-action restrictions, parameterized SQL, bounded state, loopback defaults.'),
    ('Test coverage', 'Strong regression base', '25 native CTest cases and 31 Python pytest cases passed in this review.'),
    ('Production readiness', 'Needs hardening', 'No visible CI, fuzzing, sanitizer pipeline, authentication, migration framework, or license declaration.'),
    ('Maintainability', 'Good with duplication risk', 'Clear naming and comments, but two processing paths create consistency risk.'),
])

heading(doc, 'Key Recommendations', 2)
for item in [
    'Establish CI across supported Linux and Windows toolchains with build, CTest, pytest, formatting, static analysis, and security checks.',
    'Add fuzzing and sanitizer jobs for packet decoding, TCP reassembly, rule parsing, API parsing, and PCAP handling.',
    'Define one authoritative detection/persistence pipeline or document and test the compatibility boundary between Python and C++.',
    'Add authenticated deployment guidance and explicit production controls before exposing the API beyond loopback.',
    'Declare a project license and document third-party dependency licenses before distribution.',
]: bullet(doc, item)

heading(doc, 'Scope and Method', 1)
paragraph(doc, 'This review inspected the repository source, build configuration, documentation, rules, dashboard, database models, and test suites. It included source-level architectural and security analysis, a native build, the native test suite, the Python test suite, and a Python compilation attempt. No live network capture, privileged operation, production database, or external deployment was performed.')
table(doc, ['Evidence', 'Result'], [
    ('Native source and headers', '66 files under src; approximately 6,115 lines across src/core/dashboard/database sources.'),
    ('Python source', '14 Python files, including capture, orchestration, rule engine, persistence, API proxy, and dashboard.'),
    ('Tests', '42 test files discovered; 25 C++ tests registered in CTest and 31 Python tests executed.'),
    ('Build', 'cmake --build build -j2 completed successfully.'),
    ('Native tests', '25/25 passed; 100% pass rate.'),
    ('Python tests', '31/31 passed.'),
    ('Syntax check caveat', 'A cache-writing compilation command encountered permission denied on an existing dashboard/__pycache__ artifact; this did not affect pytest or the build.'),
])

heading(doc, 'System Architecture', 1)
paragraph(doc, 'The repository contains two closely related execution surfaces. The native C++ implementation provides the reusable packet, flow, protocol, detection, behavioral, storage, API, telemetry, and platform components. The Python path provides Scapy capture/replay, rule matching, alert management, SQLite/SQLAlchemy persistence, process orchestration, and the Flask same-origin dashboard proxy.')

table(doc, ['Layer', 'Primary implementation', 'Review observations'], [
    ('Capture', 'src/capture; core/packet_capture.py', 'Supports libpcap/Npcap, PCAP replay, interface discovery, and Windows raw-socket supplementation.'),
    ('Packet processing', 'src/packet; core/packet_capture.py', 'C++ decoder handles IPv4/IPv6, VLAN, transport headers, fragments, truncation, and malformed input.'),
    ('Flows/reassembly', 'src/flow', 'Tracks bidirectional flows, TCP state, bounded capacity, idle/lifetime expiry, and reassembly outcomes.'),
    ('Protocol inspection', 'src/protocol', 'HTTP, DNS, TLS, SSH, ICMP, and generic inspectors are registered through an inspector manager.'),
    ('Detection', 'src/detection; core/detection_engine.py', 'Passive ALERT/LOG rule model, content/buffer matching, rule validation, and evidence-bearing events.'),
    ('Behavioral analytics', 'src/behavioral; core/delta_core.py', 'Port scans, host sweeps, brute-force-like failures, floods, DNS rate anomalies, and TCP anomalies.'),
    ('Persistence', 'src/storage; database/models.py; core/alert_manager.py', 'SQLite schema, parameterized filtering, bounded write queue in C++, SQLAlchemy compatibility path.'),
    ('Presentation', 'src/api; dashboard/app.py; dashboard/static', 'Native HTTP API proxied through localhost Flask dashboard with JSON export endpoints.'),
])

heading(doc, 'Evidence Flow and Boundaries', 2)
paragraph(doc, 'The documented evidence flow is coherent and security-conscious: packets are normalized before detection, stateful subsystems are bounded, alerts retain packet-derived evidence, incidents correlate on actual entity relationships, and the dashboard reports unavailable backends instead of fabricating values. Platform-specific code is intentionally constrained to interface enumeration, capture, privileges, filesystem paths, and system metrics.')
paragraph(doc, 'A notable design strength is the passive-only contract. The C++ rule parser accepts ALERT and LOG actions and explicitly rejects active actions such as DROP, REJECT, BLOCK, and REPLACE. The README and security documentation consistently state that the system does not block, inject, modify, scan, exploit, reset, or automatically respond to traffic.')

heading(doc, 'Module-by-Module Analysis', 1)
heading(doc, 'Capture and Packet Normalization', 2)
paragraph(doc, 'The Python capture layer supports live Scapy sniffing and PCAP replay, normalizes IPv4 packets, extracts transport metadata, preserves bounded payload previews, and records detailed header fields. On Windows it adds an optional raw-socket path for traffic that Npcap may not expose to the application. Duplicate suppression prevents overlapping capture paths from inflating behavioral state.')
bullet(doc, 'Strength: capture errors are isolated so the capture loop can continue, while degraded state is recorded.')
bullet(doc, 'Strength: operator-facing validation covers interface selection, database writability, privilege hints, port bounds, and executable discovery.')
bullet(doc, 'Risk: Python normalization is IPv4-focused while the native decoder supports IPv6; behavior can diverge depending on which path is used.')
bullet(doc, 'Recommendation: define a shared packet contract and add cross-path golden-vector tests for IPv4, IPv6, VLAN, fragments, TCP, UDP, and ICMP.')

heading(doc, 'Native Packet Decoder', 2)
paragraph(doc, 'The C++ decoder uses a bounded Reader abstraction with explicit availability checks before reads. It validates Ethernet, VLAN, IPv4, IPv6 extension headers, TCP header length, UDP length, and fragment handling. Invalid inputs return structured decode statuses rather than terminating the process. Telemetry increments received-packet and parser-error counters.')
bullet(doc, 'Strength: defensive parsing is visible and directly tested.')
bullet(doc, 'Risk: the decoder supports a broad protocol surface but does not appear to be protected by continuous fuzzing or sanitizer jobs.')
bullet(doc, 'Recommendation: fuzz decode() with libFuzzer/AFL++ and run AddressSanitizer/UndefinedBehaviorSanitizer in CI.')

heading(doc, 'Flow Tracking and TCP Reassembly', 2)
paragraph(doc, 'FlowManager creates stable flow identities, records client/server statistics, identifies services, updates TCP state, and expires or evicts flows using configured bounds. TCP stream trackers account for retransmissions, overlap, conflicting overlap, FIN, and RST conditions. Capacity eviction is deterministic by last-seen time and flow ID.')
bullet(doc, 'Strength: explicit maximum-flow, idle-timeout, lifetime-timeout, and reassembly configuration supports bounded memory behavior.')
bullet(doc, 'Risk: stateful flow logic is inherently sensitive to timestamp ordering, fragmentation, retransmission, and adversarial traffic patterns; sustained stress testing should be part of release qualification.')

heading(doc, 'Protocol Inspection', 2)
paragraph(doc, 'InspectorManager registers HTTP, DNS, TLS, SSH, ICMP, and generic TCP/UDP inspectors and provides reset semantics. Service identification is integrated with flow processing so detection events can carry service and protocol context.')
bullet(doc, 'Strength: inspectors are behind an abstraction and can be tested independently.')
bullet(doc, 'Recommendation: document protocol parsing limits and add corpus-based tests for malformed, truncated, encrypted, and unusual inputs.')

heading(doc, 'Signature Detection and Rules', 2)
paragraph(doc, 'The native rule parser validates JSON structure, positive SID/GID/revision identities, passive actions, protocols, ports/ranges, direction, service, buffers, content chains, metadata, severity, and thresholds. Duplicate SID/revision identities and unsupported fields are rejected. The Python DetectionEngine independently handles supported fields, content and regular-expression matching, port matching, unsupported legacy heuristic fields, alert limits, and deduplication.')
bullet(doc, 'Strength: unsupported heuristic_payload data is deliberately not treated as a payload signature, reducing false positives.')
bullet(doc, 'Strength: alert evidence includes the matched protocol and condition rather than only a generic message.')
bullet(doc, 'Risk: there are two rule/detection implementations with different feature sets and defaults; this creates a compatibility and operational drift risk.')
bullet(doc, 'Recommendation: choose a canonical engine, or publish a versioned compatibility matrix and run the same rule corpus through both paths.')

heading(doc, 'Behavioral Detection', 2)
paragraph(doc, 'Behavioral detectors implement port scans, host sweeps, repeated connection patterns, DNS query-rate anomalies, and invalid TCP flag combinations. The Python path separately implements ICMP echo and ping-sweep detection plus bounded TCP/UDP scan state. The documented scan contract correctly requires qualifying traffic, distinct destination ports, a source/destination/protocol key, a sliding window, and bounded state.')
bullet(doc, 'Strength: duplicate UDP/DNS traffic and ordinary established ACK traffic are explicitly prevented from inflating scan state.')
bullet(doc, 'Strength: port-scan evidence reports observed ports, improving analyst trust and post-event review.')
bullet(doc, 'Risk: behavioral thresholds are largely static and may require environment-specific calibration to manage false positives.')
bullet(doc, 'Recommendation: expose threshold configuration through a versioned configuration file and record detector configuration with each event.')

heading(doc, 'Storage and Incident Correlation', 2)
paragraph(doc, 'The native SQLite storage layer uses prepared statements and bound parameters, including escaped LIKE patterns. It maintains a bounded asynchronous write queue and a busy timeout for coexistence with the Python writer. The Python SQLAlchemy model defines alerts, flows, detection events, rules, incident links, incidents, statistics, and traffic logs, with lightweight additive migrations.')
bullet(doc, 'Strength: user-supplied filters are not concatenated into SQL values; search terms escape wildcard characters.')
bullet(doc, 'Strength: incident correlation requires matching source, destination, protocol, category, and time relationship rather than grouping all open alerts.')
bullet(doc, 'Risk: two persistence implementations and an ad hoc migration tuple can drift in schema or semantics over time.')
bullet(doc, 'Recommendation: introduce schema versioning/migrations, integration tests against both writers/readers, and a documented concurrency ownership model.')

heading(doc, 'API, Orchestration, and Dashboard', 2)
paragraph(doc, 'The launcher validates ports and database permissions, starts the native API and Flask dashboard, waits for readiness, optionally starts capture, and terminates child processes on shutdown. The dashboard binds to 127.0.0.1 by default and proxies GET/DELETE API operations, including JSON exports. The API surface includes status, statistics, system/configuration, alerts, traffic, incidents, flows, rules, detection events, and reset operations.')
bullet(doc, 'Strength: runtime status and stale/unavailable behavior are explicit rather than hidden behind mock data.')
bullet(doc, 'Risk: the dashboard is described as a development static host/proxy and the API has no visible authentication or authorization layer.')
bullet(doc, 'Risk: DELETE reset and data-clearing operations require deployment controls if the service is ever reachable beyond trusted localhost.')
bullet(doc, 'Recommendation: keep production deployment behind an authenticated reverse proxy, add CSRF/authorization controls for destructive operations, and document a threat model.')

heading(doc, 'Security Review', 1)
paragraph(doc, 'The repository treats packet bytes, PCAP files, rule files, and API input as untrusted data. The implementation includes several effective controls: parser bounds checks, malformed-input errors, passive action rejection, bounded API query handling, parameterized SQL, bounded flow/alert/incident/write-queue state, and localhost dashboard defaults.')
table(doc, ['Control area', 'Observed control', 'Residual concern'], [
    ('Memory safety', 'Availability checks in native decoder; bounded state in flows and behavioral detectors.', 'No fuzzing/sanitizer evidence in repository automation.'),
    ('SQL injection', 'Prepared statements and bound parameters; escaped LIKE values.', 'Schema ownership and cross-writer migrations need formalization.'),
    ('Active response', 'Only passive ALERT/LOG actions are accepted by native rule parser.', 'Deployment/network controls still matter if API is exposed.'),
    ('Input validation', 'Port, protocol, action, rule identity, query, and path checks.', 'Validation should be centralized and covered by property-based tests.'),
    ('Network exposure', 'Dashboard defaults to loopback; proxy uses local API default.', 'No authentication/authorization in reviewed dashboard/API path.'),
    ('Secrets', 'No obvious embedded credentials or shell execution paths found.', 'Add secret scanning and dependency/security scanning to CI.'),
    ('Privacy', 'Traffic metadata and evidence persist in SQLite.', 'Document retention, access permissions, and sensitive payload handling.'),
])

heading(doc, 'Testing and Verification', 1)
paragraph(doc, 'Verification during this review was positive. The existing build directory compiled successfully, all 25 registered native tests passed, and all 31 Python tests passed. The native suite covers security hardening, performance, evaluation, PCAP regression, platform providers, versioning, interfaces, capture, packet decoding, flows, reassembly, services, inspectors, buffers, rules, matching, detection, behavior, alerts, incidents, storage, API, telemetry, and API integration.')
table(doc, ['Verification command', 'Outcome'], [
    ('cmake --build build -j2', 'Passed; all targets built.'),
    ('ctest --test-dir build --output-on-failure', 'Passed; 25/25 tests.'),
    ('python3 -m pytest -q', 'Passed; 31/31 tests.'),
    ('Python compilation attempt', 'Blocked by permissions on an existing generated __pycache__ file; not a source failure.'),
])
heading(doc, 'Testing Gaps', 2)
for item in [
    'No repository-visible CI workflow was identified in the reviewed file set.',
    'No fuzzing harness or sanitizer configuration was identified.',
    'No measured live-capture performance result was produced in this review.',
    'No cross-platform build was executed on Windows.',
    'No authenticated API/deployment test was identified.',
    'No formal dependency license/SBOM workflow was identified.',
]: bullet(doc, item)

heading(doc, 'Code Quality and Maintainability', 1)
paragraph(doc, 'The code uses descriptive names, narrow classes, explicit configuration, comments around security-sensitive behavior, and test-oriented contracts. The documentation is unusually detailed for the project size and aligns well with the implementation in areas such as passive behavior, scan evidence, bounded state, and incident correlation.')
table(doc, ['Quality strength', 'Impact'], [
    ('Clear module boundaries', 'Enables focused testing and platform portability.'),
    ('Defensive comments near tricky logic', 'Improves reviewer confidence around deduplication, scan qualification, and SQL binding.'),
    ('Deterministic tests', 'Makes regressions in protocol, behavior, persistence, and API semantics easier to identify.'),
    ('Operational README', 'Reduces setup and database-permission errors for operators.'),
])
paragraph(doc, 'The main maintainability concern is duplication: Python and C++ both implement capture-adjacent processing, rule interpretation, behavioral logic, and persistence. Duplication can be acceptable as a compatibility layer, but it should be made explicit with shared fixtures, schema contracts, and release tests that prevent semantic drift.')

heading(doc, 'Prioritized Roadmap', 1)
table(doc, ['Priority', 'Action', 'Why it matters', 'Suggested horizon'], [
    ('P0', 'Add CI for Linux and Windows build/test paths.', 'Prevents regressions in the cross-platform contract and makes quality repeatable.', 'Immediate'),
    ('P0', 'Add fuzzing and ASan/UBSan jobs for parsers and reassembly.', 'IDS components process adversarial bytes and are high-value targets.', 'Immediate'),
    ('P0', 'Declare a license and dependency licensing policy.', 'Required for responsible distribution and downstream adoption.', 'Immediate'),
    ('P1', 'Make one detection/persistence path authoritative or define compatibility tests.', 'Reduces divergent alerts, schemas, and operational behavior.', 'Near term'),
    ('P1', 'Formalize SQLite schema migrations and writer concurrency.', 'Protects upgrades and mixed native/Python operation.', 'Near term'),
    ('P1', 'Add API authentication/authorization guidance and destructive-operation controls.', 'Prevents accidental or unauthorized data clearing when deployed beyond localhost.', 'Near term'),
    ('P2', 'Add protocol corpora, property-based tests, and sustained replay benchmarks.', 'Improves confidence against malformed input and realistic workloads.', 'Medium term'),
    ('P2', 'Add retention/privacy controls for stored traffic evidence.', 'Reduces operational and data-protection risk.', 'Medium term'),
])

heading(doc, 'Conclusion', 1)
paragraph(doc, 'Delta-NIDS has a credible engineering foundation: the architecture is understandable, the passive-only security boundary is explicit, parsing and storage code show defensive intent, and the current automated suites pass cleanly. The project should be treated as a well-tested early-stage IDS rather than a fully production-hardened security appliance. The highest-value next step is to convert the existing good engineering practices into repeatable release controls—cross-platform CI, fuzzing/sanitizers, canonical pipeline contracts, formal migrations, deployment security, and licensing.')

heading(doc, 'Appendix A — Repository Snapshot', 1)
table(doc, ['Item', 'Value'], [
    ('Repository', 'git@github.com:DavidKuriyan/delta-ids.git'),
    ('Branch', 'main'),
    ('Reviewed commit', '0b9b31e — Fix dashboard search and restart state handling'),
    ('Native language', 'C++17'),
    ('Python stack', 'Python, Scapy, Flask, SQLAlchemy, SQLite'),
    ('Native source/header files', '66'),
    ('Python source files', '14'),
    ('Test files discovered', '42'),
    ('Native tests passed', '25'),
    ('Python tests passed', '31'),
    ('Declared license', 'None currently declared in repository'),
])

heading(doc, 'Appendix B — Reproduction Commands', 1)
for cmd in [
    'cmake --build build -j2',
    'ctest --test-dir build --output-on-failure',
    'python3 -m pytest -q',
    'cmake -S . -B build -DDELTA_NIDS_BUILD_TESTS=ON',
]:
    p = doc.add_paragraph()
    r = p.add_run(cmd); r.font.name = 'Consolas'; r.font.size = Pt(9); r.font.color.rgb = RGBColor(60,60,60)

# Core properties for document metadata
props = doc.core_properties
props.title = 'Delta-NIDS Full Codebase Analysis Report'
props.subject = 'Architecture, security, quality, testing, and recommendations'
props.author = 'Codebuff Engineering Review'
props.keywords = 'Delta-NIDS, IDS, code review, security, architecture'
props.comments = 'Generated from repository source and verification results.'

doc.save(OUT)
print(OUT)
